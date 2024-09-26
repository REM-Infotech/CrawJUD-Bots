from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import os
from typing import Type
import shutil
from docx.shared import Pt
from docx.text.font import Font

from time import sleep

def replace_placeholder(output_path, replacements: dict, pid: str) -> str:
    # Abra o documento
    
    namefile = f"MANIFESTAÇÃO - CUMPRIMENTO OBRIGAÇÃO DE PAGAR {replacements.get('[[NUMERO_PROCESSO]]')} - {pid}.docx"
    output_doc = os.path.join(output_path, "protocolar", namefile)
    doc_path = os.path.join(os.getcwd(), "docs", "esqueleto.docx") 
    doc = Document(doc_path)

    # Itere sobre todos os parágrafos do documento
    for paragraph in doc.paragraphs:
        for placeholder, content in replacements.items():
            index = 0
            if placeholder in paragraph.text:
                # Encontre a posição do placeholder
                position = paragraph.text.find(placeholder)
                # Divida o texto ao redor do placeholder
                before_text = paragraph.text[:position]
                after_text = paragraph.text[position + len(placeholder):]
                
                # Salve as propriedades de formatação do parágrafo original
                original_font_name = paragraph.runs[0].font
                original_font_name: Type[Font] = original_font_name.name
                
                original_font_size = Pt(12)

                # Limpe o parágrafo
                paragraph.clear()
                
                # Adicione o texto antes do placeholder
                run = paragraph.add_run(before_text)
                run.bold = True
                run.font.size = original_font_size
                
                # Adicione a imagem ou texto em negrito
                if any(extension in content for extension in [".png", ".jpg", "jpeg"]):
                    run = paragraph.add_run()
                    run.add_picture(content, width=Inches(6))  # Ajuste o tamanho conforme necessário
                    run.add_break(WD_BREAK.LINE)
                    run.add_break(WD_BREAK.LINE)
                    run = paragraph.add_run(after_text)
                    run.bold = True
                else:
                    # Adicione o conteúdo em negrito
                    run = paragraph.add_run(content)
                    run.bold = True
                    run.font.name = original_font_name
                    run.font.size = original_font_size
                    # Adicione o texto restante
                    run = paragraph.add_run(after_text)
                    run.bold = True
                
                index = index + 1

    # Salve o documento modificado
    doc.save(output_doc)
    shutil.copy(output_doc, os.path.join(output_path, replacements.get('[[NUMERO_PROCESSO]]'), namefile))
    
    return namefile